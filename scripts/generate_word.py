import os
import sys
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def set_run_style(run, size=11, bold=False, italic=False, color=(0,0,0), font_name='맑은 고딕'):
    run.font.name = font_name
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = RGBColor(*color)

def add_styled_paragraph(doc, text, level=0, is_bullet=False):
    if level == 1: # Title #
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_style(run, size=20, bold=True, color=(0, 51, 102))
        p.space_after = Pt(18)
    elif level == 2: # Header ##
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_style(run, size=14, bold=True, color=(51, 102, 153))
        p.space_before = Pt(12)
        p.space_after = Pt(6)
    elif level == 3: # Subheader ###
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_style(run, size=12, bold=True, color=(102, 102, 102))
    elif is_bullet:
        p = doc.add_paragraph(style='List Bullet')
        run = p.add_run(text)
        set_run_style(run, size=11)
    else:
        p = doc.add_paragraph()
        run = p.add_run(text)
        set_run_style(run, size=11)
    return p

def convert_md_to_docx(md_path: Path, output_path: Path, job_id: str):
    if not md_path.exists():
        print(f"Error: MD path {md_path} not found.")
        return False

    doc = Document()
    md_content = md_path.read_text(encoding="utf-8")
    
    # Simple Title Page
    title_para = doc.add_paragraph()
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_para.space_before = Inches(2)
    run = title_para.add_run("STRATEGIC DEEP ANALYSIS REPORT\n")
    set_run_style(run, size=24, bold=True, color=(0, 51, 102))
    
    info_para = doc.add_paragraph()
    info_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = info_para.add_run(f"Project ID: {job_id}\nDate: {datetime.now().strftime('%Y-%m-%d')}")
    set_run_style(run, size=12, color=(128, 128, 128))
    
    doc.add_page_break()

    # Parse MD Lines
    lines = md_content.split('\n')
    for line in lines:
        line = line.strip()
        if not line: continue
        
        # Headers
        if line.startswith('# '):
            add_styled_paragraph(doc, line[2:], level=1)
        elif line.startswith('## '):
            add_styled_paragraph(doc, line[3:], level=2)
        elif line.startswith('### '):
            add_styled_paragraph(doc, line[4:], level=3)
        # Bullets
        elif line.startswith('- ') or line.startswith('* '):
            # Remove bold markdown **...**
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line[2:])
            add_styled_paragraph(doc, clean_text, is_bullet=True)
        # Regular text
        else:
            clean_text = re.sub(r'\*\*(.*?)\*\*', r'\1', line)
            add_styled_paragraph(doc, clean_text)

    # Footer
    for section in doc.sections:
        footer = section.footer
        fp = footer.paragraphs[0]
        fp.text = f"© 2026 Antigravity | Standard Strategic MD Format"
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(str(output_path))
    return True

def main():
    if len(sys.argv) < 2:
        print("Usage: python generate_word.py <job_id>")
        sys.exit(1)
        
    job_id = sys.argv[1]
    root = Path(__file__).resolve().parents[1]
    md_path = root / "outputs" / "02_deep_analysis" / f"{job_id}_deep_report.md"
    output_dir = root / "outputs" / "03_final_reports"
    output_dir.mkdir(parents=True, exist_ok=True)
    output_file = output_dir / f"{job_id}_depth_report.docx"
    
    print(f"[Word] Converting {md_path.name} to professional Word doc...")
    if convert_md_to_docx(md_path, output_file, job_id):
        print(f"[Word] Success: Generated {output_file.name}")
    else:
        sys.exit(1)

if __name__ == "__main__":
    main()
