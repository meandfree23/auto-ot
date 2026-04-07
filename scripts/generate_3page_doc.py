import sys
import os
import re
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def generate_3page_doc(job_id):
    """
    Reconstruct the 3-page Word strategy summary from deep analysis markdown.
    Structure:
    Page 1: Market Intelligence & Context
    Page 2: Strategic Proposals & Key Solutions
    Page 3: Creative Drafts & Implementation Roadmap
    """
    ROOT = Path(__file__).resolve().parents[1]
    input_file = ROOT / "outputs" / "02_deep_analysis" / f"{job_id}_deep_report.md"
    output_file = ROOT / "outputs" / "03_final_reports" / f"{job_id}_exec_summary_3p.docx"

    if not input_file.exists():
        print(f"Error: Input file {input_file} not found.")
        return False

    with open(input_file, "r", encoding="utf-8") as f:
        content = f.read()

    doc = Document()

    # Style: Standard professional
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(10)

    def add_section_title(text):
        p = doc.add_paragraph()
        run = p.add_run(text)
        run.bold = True
        run.font.size = Pt(14)
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        doc.add_paragraph() # Spacer

    # Helper to clean markdown formatting
    def clean_md(text):
        text = re.sub(r'#+\s*', '', text)
        text = re.sub(r'\*\*(.*?)\*\*', r'\1', text)
        return text.strip()

    # Split content by H1 or H2
    sections = re.split(r'\n(?=#+ )', content)
    
    # 📄 PAGE 1: Market Intelligence
    add_section_title("PAGE 1: MARKET INTELLIGENCE & CONTEXT")
    for section in sections[:3]: # Grab first few sections
        lines = section.strip().split('\n')
        if not lines: continue
        title = clean_md(lines[0])
        body = "\n".join(lines[1:])
        p = doc.add_paragraph()
        p.add_run(f"■ {title}").bold = True
        doc.add_paragraph(clean_md(body))
    
    doc.add_page_break()

    # 📄 PAGE 2: Strategic Proposals
    add_section_title("PAGE 2: STRATEGIC PROPOSALS & SOLUTIONS")
    for section in sections[3:6]:
        lines = section.strip().split('\n')
        if not lines: continue
        title = clean_md(lines[0])
        body = "\n".join(lines[1:])
        p = doc.add_paragraph()
        p.add_run(f"▶ {title}").bold = True
        doc.add_paragraph(clean_md(body))

    doc.add_page_break()

    # 📄 PAGE 3: Creative & Roadmap
    add_section_title("PAGE 3: CREATIVE DRAFTS & ROADMAP")
    for section in sections[6:]:
        lines = section.strip().split('\n')
        if not lines: continue
        title = clean_md(lines[0])
        body = "\n".join(lines[1:])
        p = doc.add_paragraph()
        p.add_run(f"✦ {title}").bold = True
        doc.add_paragraph(clean_md(body))

    doc.save(output_file)
    print(f"Success: {output_file}")
    return True

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python3 generate_3page_doc.py <job_id>")
        sys.exit(1)
    generate_3page_doc(sys.argv[1])
